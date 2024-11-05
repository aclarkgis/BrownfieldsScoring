import csv
import smtplib
from flask import Flask, render_template, request
from docx import Document
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import os

app = Flask(__name__)

# Sample data structure (questions and groups)
groups = [
     {
        "name": "Ownership",
        "description": "Ownership (or a path to ownership) is a critical component to a realistic assessment of the redevelopment potential of a site. Obtaining site control can take years because of impediments such as liens, bankruptcy, abandonment, or uncooperative and unrealistic owners.",
        "allow_multiple": True,
        "questions": [
            {"text": "Do you own the property?", "value": 10},
            {"text": "Can you get site control of the property through a sales contract?", "value": 6, "id": "group1_question1"},
            {"text": "Can you get site control by clearing or foreclosing on a lien?", "value": 2, "id": "group1_question2"},
            {"text": "Do you still need to determine a path to site control?", "value": 0, "id": "group1_question3"}
        ]
    },
    {
        "name": "Site Use",
        "description": "The anticipated benefits of site reuse are most often the impetus for brownfield clean-up and redevelopment. Reuse opportunities drive clean-up requirements, project economics and financing options and community impacts. The intended reuse of a property should bring the market together with the community's goals and aspirations.",
        "allow_multiple": True,
        "questions": [
            {"text": "Do you have an interested, qualified user?", "value": 20, "id": "group2_question1"},
            {"text": "Has anyone expressed interest in reusing the site in the past year?", "value": 8, "id": "group2_question2"},
            {"text": "Have any properties adjacent to the site been redeveloped within the past two years?", "value": 7, "id": "group2_question3"},
            {"text": "Have any properties within ½ mile of the site been redeveloped within the past two years?", "value": 5, "id": "group2_question4"}
        ],
        "subgroups": [
            {
                "name": "Consistency",
                "description": "",
                "allow_multiple": False,
                "questions": [
                    {"text": "Is it complementary or consistent with surrounding uses?", "value": 2, "id": "group2_question5"},
                    {"text": "Is it consistent with the zoning on the property?", "value": 2, "id": "group2_question6"},
                    {"text": "Is it consistent with the community's plans and goals?", "value": 2, "id": "group2_question7"},
                    {"text": "You don't know?", "value": 0, "id": "group2_question8"}
                ]
            },
            {
                "name": "Other Properties",
                "description": "",
                "allow_multiple": False,
                "questions": [
                    {"text": "Are there other existing community assets nearby that make this site desirable?", "value": 4, "id": "group2_question9"},
                    {"text": "Are there other planned community assets that attract people or investment that could help make this site desirable in the near future?", "value": 2, "id": "group2_question10"},
                    {"text": "Are there temporary uses for the site that could improve its long-term marketability?", "value": 1, "id": "group2_question11"},
                    {"text": "Are adjacent or nearby properties vacant or blighted?", "value": 0, "id": "group2_question12"},
                    {"text": "You don't know?", "value": 0, "id": "group2_question13"}
                ]
            }
        ]
    },
    {
        "name": "Land Characteristics",
        "description": "The size of a property (or multiple properties that can be assembled together into one property) contributes to the redevelopment potential of each brownfield site. There are more potential users for mid-sized sites than for very large or very small sites.",
        "allow_multiple": True,
        "questions": [],
        "subgroups": [
            {
                "name": "Size",
                "description": "The size of a property contributes to the redevelopment potential.",
                "allow_multiple": True,
                "questions": [
                    {"text": "Is the property itself between 1 and 5 acres?", "value": 10, "id": "group3_question1"},
                    {"text": "Can the property be assembled with others to be 1-5 acres?", "value": 8, "id": "group3_question2"},
                    {"text": "Is the property between 5 and 15 acres?", "value": 8, "id": "group3_question3"},
                    {"text": "Can the property be assembled with others to be 5-15 acres?", "value": 6, "id": "group3_question4"},
                    {"text": "Is the property itself between 15 and 50 acres?", "value": 6, "id": "group3_question5"},
                    {"text": "Can the property be assembled with others to be 15-50 acres?", "value": 8, "id": "group3_question6"},
                    {"text": "Is the property over 50 acres?", "value": 4, "id": "group3_question7"},
                    {"text": "Is the property between ½ and 1 acre?", "value": 2, "id": "group3_question8"},
                    {"text": "Is the property less than 1/2 acre?", "value": 0, "id": "group3_question9"},
                    {"text": "Is the property itself between 15 and 50 acres?", "value": 6, "id": "group3_question10"}
                ]
            },
            {
                "name": "Flooding",
                "description": "Flooding can increase redevelopment costs.",
                "allow_multiple": True,
                "questions": [
                    {"text": "There are no obvious or known flooding problems on the site?", "value": 5, "id": "group3_question11"},
                    {"text": "There are known flooding problems, but green infrastructure is an option?", "value": 2, "id": "group3_question12"},
                    {"text": "There are known flooding problems and green infrastructure is not an option?", "value": 0, "id": "group3_question13"}
                ]
            },
            {
                "name": "Soil",
                "description": "Soil conditions can make building on a site more costly.",
                "allow_multiple": True,
                "questions": [
                    {"text": "Are there no known problems with the soil conditions or compaction on the site?", "value": 5, "id": "group3_question14"},
                    {"text": "Are there known problems with the soil conditions or compaction on the site?", "value": 0, "id": "group3_question15"}
                ]
            }
        ]
    },
    {
        "name": "Community Characteristics",
        "description": "The condition of adjacent properties impacts the marketability of a site.",
        "allow_multiple": True,
        "questions": [],
        "subgroups": [
            {
                "name": "Actions",
                "description": "In the past two years, have adjacent or nearby properties been...",
                "allow_multiple": True,
                "questions": [
                    {"text": "Improved?", "value": 10, "id": "group4_question1"},
                    {"text": "Maintained?", "value": 6, "id": "group4_question2"},
                    {"text": "Blighted?", "value": 0, "id": "group4_question3"},
                    {"text": "A mixture of the above?", "value": 3, "id": "group4_question4"}
                ]
            },
            {
                "name": "Planned Investment and Nearby Uses",
                "description": "Planned investments and nearby uses can impact the marketability of a site.",
                "allow_multiple": True,
                "questions": [
                    {"text": "Is there a large planned investment or use nearby that provides untapped market opportunities?", "value": 10, "id": "group4_question5"},
                    {"text": "Are there planned or existing smaller uses nearby that will bring more investment?", "value": 6, "id": "group4_question6"},
                    {"text": "Are there smaller uses that will discourage investment?", "value": 0, "id": "group4_question7"},
                    {"text": "Is there a large use nearby that will discourage investment?", "value": 0, "id": "group4_question8"},
                    {"text": "Nearby uses will not impact this site.", "value": 0, "id": "group4_question9"}
                ]
            },
            {
                "name": "Crime and Safety",
                "description": "The perception that a site is in a safe location will impact the demand and potential uses for a site. The ability to get to and from a site safely can be as important as the safety of the actual site.",
                "allow_multiple": True,
                "questions": [
                    {"text": "Is the crime rate where your property is located lower than average for the city or town? ", "value": 8, "id": "group4_question10"},
                    {"text": " Is the crime rate where your property is located about average the city or town?", "value": 4, "id": "group4_question11"},
                    {"text": "Is the crime rate where your property is located worse than average for your city or town?", "value": 0, "id": "group4_question12"}
                ]
            }
        ]
    },
    {
        "name": "Community Capacity",
        "description": "Local government can facilitate or inhibit brownfield redevelopment in many ways. Brownfield transactions tend to be complicated and government help is often needed to overcome these complications. A community that is capable of and willing to provide this help can make an otherwise impossible situation workable.",
        "allow_multiple": True,
        "questions": [
           {"text": "Does your community have a successful brownfield redevelopment program through which other properties have been redeveloped in the past two years?", "value": 6, "id": "group5_question1"},
           {"text": "Is this community lacking a successful brownfield program? ", "value": 0, "id": "group5_question2"},
           {"text": "You don't know", "value": 0, "id": "group5_question3"}
        ],
        "subgroups" : [
            {
                "name": "Community Plan",
                "description": "Community consensus about the future use of a brownfield site removes a significant amount of risk facing the potential developer of a site. Protracted battles about zoning and environmental impacts can costs, uncertainty and even months or years to a redevelopment project.",
                "allow_multiple": True,
                "questions": [
                    {"text": "Does your community have a generally accepted redevelopment plan less than five years old that includes this property? ", "value": 6, "id": "group5_question4"},
                    {"text": "This community has no redevelopment plan that includes this property", "value": 0, "id": "group5_question5"},
                    {"text": "Is there agreement about the reuse of this property? ", "value": 0, "id": "group5_question6"},
                    {"text": "Is there controversy about the redevelopment goals for this property?", "value": 0, "id": "group5_question7"},
                    {"text": "You dont know?", "value": 0, "id": "group5_question8"}
                ]    
            }
        ]
    },
    {
        "name": "Redevelopment Incentives",
        "description": "Some brownfield sites are up-side-down, that is the cost of cleaning them up and redeveloping them are greater than the economic value of the redevelopment to the developer. Other sites come close to this point. In such cases redevelopment incentives are needed to pay these legacy costs so development can be economically viable. Many communities find it reasonable to provide incentives because they benefit from redevelopment through impacts such as new tax revenue, job creation, services, and increased property values.",
        "allow_multiple": True,
        "questions": [
           {"text": "Is it included in a formally designated redevelopment area or Tax Increment Financing (TIF) District?", "value": 8, "id": "group6_question1"},
           {"text": "Is it eligible for Historic or New Markets Tax Credits?", "value": 6, "id": "group6_question2"},
           {"text": "Is it eligible for other redevelopment grants or loans", "value": 8, "id": "group6_question3"},
           {"text": "Are financial incentives lacking for this property? ", "value": 0, "id": "group6_question4"},
           {"text": "You don't know", "value": 0, "id": "group6_question5"}
        ]
    },
    {
        "name": "Infrastructure Amenities",
        "description": "",
        "allow_multiple": True,
        "questions": [],
        "subgroups" : [
            {
            "name": "Public Transportation",
            "description": "Access to a site impacts its marketability and in many areas public transit is a key form of transportation for potential users (workers, students, residents, shoppers, etc.) More users will find easily accessible sites attractive thus creating a stronger market for these sites.",
            "allow_multiple": True,
            "questions": [
                {"text": "Is this property is located within 1/2 mile of a public transit train stop?", "value": 8, "id": "group7_question1"},
                {"text": "Is this property is located within 1/2 mile of a public transit bus stop?", "value": 6, "id": "group7_question2"},
                {"text": "Is public transit access lacking within 1/2 mile if the property? ", "value": 0, "id": "group7_question3"}
                ]  
            },
            {
                "name": "Roads",
                "description": "Road access also impacts the marketability of a site. Most people travel by car and most goods travel by truck making highway and arterial street access a very important asset. In some cases, such as retail uses, visibility and the amount of traffic traveling past a site each day play key roles in impacting desirability.",
                "allow_multiple": True,
                "questions": [
                    ]  
            },
            {
                "name": "Water & Sewer",
                "description": "If water and sewer service are not adequate they can add significant costs to site redevelopment.",
                "allow_multiple": True,
                "questions": [
                    {"text": "Does adequate water/sewer come to the site?", "value": 5, "id": "group7_question4"},
                    {"text": "(s adequate water/sewer missing from the site?", "value": 0, "id": "group7_question5"}
                    ]  
            },
            {
                "name": "Electricity",
                "description": "If electrical service is not adequate it can add significant costs to site redevelopment.",
                "allow_multiple": True,
                "questions": [
                    {"text": "Does adequate electricity come to the site?", "value": 5, "id": "group7_question6"},
                    {"text": "Is adequate electricity missing from the site?", "value": 0, "id": "group7_question7"}
                    ]  
            },
            {
                "name": "Heating Fuel",
                "description": "If heating fuel is not accessible at the site, in order to remedy this problem, significant costs can be added to site redevelopment.",
                "allow_multiple": True,
                "questions": [
                    {"text": "Does adequate heating fuel come to the site?", "value": 5, "id": "group7_question8"},
                    {"text": "Is adequate heating fuel missing from the site?", "value": 0, "id": "group7_question9"}
                    ]  
            },
            {
                "name": "Internet",
                "description": "Many potential users require high speed internet access. Existing access can be a big asset for a site. Lack of adequate service can be added costs or even make a project not viable at a particular location.",
                "allow_multiple": True,
                "questions": [
                    {"text": "Does the site have adequate internet access", "value": 5, "id": "group7_question10"},
                    {"text": "Is the site lacking adequate internet access", "value": 0, "id": "group7_question11"}
                    ]  
            }
        ]
    },
    {
      "name": "Environmental Conditions",
        "description": "Knowledge about the environmental conditions of a site help potential users determine whether they are interested in a site or not. It costs money to obtain this information. It is easier and cheaper for potential users to evaluate a site when this information is readily available.",
        "allow_multiple": True,
        "questions": [
            {"text": "Does the site have an NFR", "value": 10, "id": "group8_question1"},
            {"text": "Does the owner have recent reports that indicate what clean-up is needed", "value": 8, "id": "group8_question2"}
            ],
            "subgroups" : [
            {
            "name": "Contamination",
            "description": "Does the owner have recent reports that indicate what additional environmental investigation is needed? Choose one.",
            "allow_multiple": True,
            "questions": [
                {"text": "Unlikely Contaminated", "value": 8, "id": "group8_question3"},
                {"text": "Likely Contaminated - Light", "value": 4, "id": "group8_question4"},
                {"text": "Likely Contaminated - Moderate", "value": 0, "id": "group8_question5"},
                {"text": "Likely Contaminated - Substantial", "value": 0, "id": "group8_question6"},
                {"text": "Unknown -- If there are no reports, based on the history of the property do you suspect that environmental clean-up is needed? ", "value": 0, "id": "group8_question7"}
                ]  
            },
            {
            "name": "Environmental Investigation Resources",
            "description": "Environmental investigation can be expensive. Potential users are often reluctant to pay money to determine the environmental problems on a site they do not own-- especially if alternative sites are available that do not require such investigation. Often, a source of funds to characterize the environmental problem can encourage potential users to continue evaluating a site for redevelopment.",
            "allow_multiple": True,
            "questions": [
                {"text": "Is there an existing source of funds to conduct the environmental investigations needed?", "value": 6, "id": "group8_question7"},
                {"text": "Is there a potential source of funds to conduct the environmental investigations needed? ", "value": 2, "id": "group8_question8"},
                {"text": "Are sources of funds lacking to help conduct environmental investigations?", "value": 0, "id": "group8_question9"},
                {"text": "You don't know", "value": 0, "id": "group8_question10"}
                ]  
            },
            {
            "name": "Environmental Remediation Costs",
            "description": "Environmental remediation costs can have a big impact on the redevelopment potential of a site. The smaller the remediation costs in relation to the value of the property the more likely a developer will be able to pay these costs and still make a profit without the hassle and uncertainty applying for government incentives to cover them.",
            "allow_multiple": True,
            "questions": [
                {"text": "Do you suspect that the cost of clean-up (including demolition) is as less than 50% of the value of the property?", "value": 7, "id": "group8_question11"},
                {"text": "Do you suspect that the cost of clean-up (including demolition) is as between 50% and 100% of the value of the property?", "value": 3, "id": "group8_question12"},
                {"text": "Do you suspect that the cost of clean-up (including demolition if needed) will be greater than the value of the property?", "value": 0, "id": "group8_question13"},
                {"text": "You don't know", "value": 0, "id": "group8_question14"}
                ]  
            },
            {
            "name": "Environmental Remediation Resources",
            "description": "Sometimes there is no economically viable option to redevelop a site without government incentives. Government incentives can take a long time to access. If they are already in place they can make the process to obtain them easier, less risky and quicker.",
            "allow_multiple": True,
            "questions": [
                {"text": "Is there an existing source of funds to help pay the remediation and/or demolition costs if needed? ", "value": 7, "id": "group8_question15"},
                {"text": "Is there a potential source of funds to help pay the remediation and/or demolition costs if needed? ", "value": 7, "id": "group8_question16"},
                {"text": "Are sources of funds lacking to conduct environmental remediation or demolition?  ", "value": 0, "id": "group8_question17"},
                {"text": "You don't know", "value": 0, "id": "group8_question18"}
                ]  
            }
        ]  
    },
    {
        "name": "Building Characteristics",
        "description": "Most end users require a building of some sort. Sites with existing buildings can be appealing to specific users that require a similar structure. The cost and time involved in new construction can deter many potential users that are interested in an easier, cheaper alternative. Also, most users do not want to build a facility that costs significantly more than the value of other buildings in the general area.",
        "allow_multiple": True,
        "questions": [
            {"text": "Does the property have a substantial building on it? ", "value": 0, "id": "group9_question1"},
            {"text": "If the property just land (without a substantial building) is the cost of new construction in the area is less than 30% more than the sale price of desirable buildings?", "value": 8, "id": "group9_question2"},
            {"text": "If the property just vacant land (without a substantial building on it) is the cost of new building construction in the area is over 30% more than the sale price of desirable buildings? ", "value": 0, "id": "group9_question3"}
            ],
    },
    {
        "name": "Building Quality",
        "description": "An existing building can be an asset as described above. But, it can also be a detriment if it needs too much investment in order to be useful or if it is not usable and needs to be demolished.",
        "allow_multiple": True,
        "questions": [
            {"text": "Does the property have a substantial usable building on it?", "value": 7, "id": "group9_question4"},
            {"text": "Does the property have a substantial building that can be usable with modest effort such as cosmetics and demising walls? ", "value": 6, "id": "group9_question4"},
            {"text": "Does the property have a substantial building that needs significant work done to one major item (such as a roof, windows, HVAC system) to be usable", "value": 2.5, "id": "group9_question5"},
            {"text": "Does the property have a substantial building that needs significant work done to more than one major item (such as a roof, windows, HVAC system) to be usable", "value": 0, "id": "group9_question6"},
            {"text": "Does the property have a substantial building that needs to be demolished? ", "value": 0, "id": "group9_question7"}
            ],
    },
        {
        "name": "Building Stories",
        "description": "The market for modern single story buildings that can be used as is is greater than the market for multi-story buildings that require adaptive reuse.",
        "allow_multiple": True,
        "questions": [
            {"text": "Is the building mostly single story?", "value": 5, "id": "group9_question8"},
            {"text": "Is the building mostly 2-3 stories?", "value": 1.5, "id": "group9_question9"},
            {"text": "Is the building 4 + stories?", "value": 0, "id": "group9_question10"}
        ]
    },
    {
        "name": "Building Size",
        "description": "There are more potential users for mid-sized buildings than for very large or very small buildings.",
        "allow_multiple": True,
        "questions": [
            {"text": "Is the building footprint between 5,000 and 25,000 square feet? ", "value": 4, "id": "group9_question11"},
            {"text": "Is the building footprint between 25,000 and 150,000 square feet? ", "value": 2.5, "id": "group9_question12"},
            {"text": "Is the building footprint over 150,000 square feet? ", "value": 1, "id": "group9_question13"}
        ]
    },
    {
        "name": "Building Flexibility",
        "description": "Some buildings were built to be more flexible space that can be reused in many ways by many potential users. Other buildings were designed for a specific use and are not easily reusable by many potential users.",
        "allow_multiple": True,
        "questions": [
            {"text": "Is the building set up best for just one user?", "value": 4, "id": "group9_question14"},
            {"text": "Is the building easily divisible for multiple users?", "value": 3, "id": "group9_question15"},
            {"text": "Is the building divisible with moderate investment?", "value": 0, "id": "group9_question16"},
            {"text": "Do you still need to figure this out?", "value": 0, "id": "group9_question17"}
        ]
    },
    {
        "name": "Building Floor Area Ratio (FAR)",
        "description": "Buildings that take up a smaller footprint on a site leave more flexibility for potential users of that site. Potential users may need room for truck access, parking, expansion, outdoor storage, detention ponds, etc. to operate effectively. Thus, buildings that take up a large portion of a site do not provide the flexibility that many potential users need. Note: 1 acre = 43,000 square feet. Floor Area Ratio (FAR) = (Building Footprint in square feet) × (# Stories) ÷ Acreage in square feet",
        "allow_multiple": True,
        "questions": [
            {"text": "Is the Floor Area Ratio .3 or less? Bldg floor sq ft/plot sq ft", "value": 5, "id": "group9_question18"},
            {"text": "Is the Floor Area Ratio between .3 and 1? Bldg floor sq ft/plot sq ft", "value": 3, "id": "group9_question19"},
            {"text": " Is the Floor Area Ratio greater than 1? Bldg floor sq ft/plot sq ft", "value": 0, "id": "group9_question20"}
        ]
    }
    # your group data here...
]

# Function to generate the form
def get_form():
    return groups

@app.route('/')
def index():
    return render_template('index.html', groups=get_form())

@app.route('/submit', methods=['POST'])
def submit_form():
    project_name = request.form.get('project_name', 'Unnamed Project')
    recipient_email = request.form.get('email', '')

    total_score = 0
    group_scores = {}
    answers = {}  # Store the answers for each question

    for group in get_form():
        group_score = 0

        if group['allow_multiple']:
            selected_values = request.form.getlist(f'question_{group["name"]}[]')
            group_score += sum(int(value) for value in selected_values)
            answers[group['name']] = selected_values  # Store selected values for this group
        else:
            selected_value = request.form.get(f'question_{group["name"]}')
            if selected_value:
                group_score += int(selected_value)
                answers[group['name']] = selected_value  # Store the selected value

        if 'subgroups' in group:
            for subgroup in group['subgroups']:
                subgroup_score = 0
                if subgroup['allow_multiple']:
                    selected_values = request.form.getlist(f'question_{subgroup["name"]}[]')
                    subgroup_score += sum(int(value) for value in selected_values)
                    answers[subgroup['name']] = selected_values  # Store selected values for this subgroup
                else:
                    selected_value = request.form.get(f'question_{subgroup["name"]}')
                    if selected_value:
                        subgroup_score += int(selected_value)
                        answers[subgroup['name']] = selected_value  # Store the selected value

                group_score += subgroup_score

        group_scores[group['name']] = group_score
        total_score += group_score

    # Save submission to CSV file
    save_to_csv(project_name, total_score, group_scores, answers)

    # Generate the Word document with answers
    doc_path = generate_word_doc(project_name, total_score, group_scores, answers)

    # Send the Word document via email (comment this out to disable emailing)
    # if recipient_email:
    #     send_email_with_attachment(recipient_email, project_name, doc_path)

    # Return result page displaying group scores and total score
    result = f"<h2>Project: {project_name}</h2>"
    for group_name, score in group_scores.items():
        result += f"<p>{group_name} Score: {score}</p>"
    result += f"<p>Total Score: {total_score}</p>"

    return result


def save_to_csv(project_name, total_score, group_scores, answers):
    """Save the project submission to a CSV file"""
    file_exists = False
    file_path = 'submissions.csv'
    
    try:
        # Check if the file exists
        with open(file_path, 'r', newline='') as csvfile:
            file_exists = True
    except FileNotFoundError:
        # File doesn't exist yet, we will create it
        pass

    # Open the file for appending new data
    with open(file_path, 'a', newline='') as csvfile:
        fieldnames = ['Project Name', 'Total Score'] + list(group_scores.keys()) + list(answers.keys())
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)

        # Write header only if the file didn't exist before
        if not file_exists:
            writer.writeheader()

        # Create the row to save
        row = {'Project Name': project_name, 'Total Score': total_score}
        row.update({k: v for k, v in group_scores.items()})  # Add group scores
        row.update({k: ', '.join(map(str, v)) if isinstance(v, list) else v for k, v in answers.items()})  # Add answers

        # Write the row to the CSV file
        writer.writerow(row)

def generate_word_doc(project_name, total_score, group_scores, answers):
    """Generate a Word document with the project data"""
    doc = Document()
    doc.add_heading(f'Project: {project_name}', 0)

    # Add scores for each group
    for group_name, score in group_scores.items():
        doc.add_paragraph(f'{group_name} Score: {score}')

    doc.add_paragraph(f'Total Score: {total_score}')

    # Add answers for each question
    doc.add_heading('Answers:', level=1)
    for question, answer in answers.items():
        if isinstance(answer, list):
            doc.add_paragraph(f'{question}: {", ".join(answer)}')  # Join multiple answers
        else:
            doc.add_paragraph(f'{question}: {answer}')  # Single answer

    # Save the document
    doc_path = f"{project_name.replace(' ', '_')}_submission.docx"
    doc.save(doc_path)

    return doc_path

def send_email_with_attachment(recipient_email, project_name, file_path):
    """Send an email with the Word document attached"""
    sender_email = 'test'  # Replace with your email
    sender_password = 'your_password'  # Replace with your password
    subject = f'Submission Results for {project_name}'

    # Email message
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = recipient_email
    msg['Subject'] = subject
    msg.attach(MIMEText(f'Attached are the results for your project: {project_name}', 'plain'))

    # Attach the file
    with open(file_path, 'rb') as file:
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(file.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename={os.path.basename(file_path)}')
        msg.attach(part)

    # Send the email
    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender_email, sender_password)
        server.sendmail(sender_email, recipient_email, msg.as_string())
        server.quit()
        print(f'Email sent to {recipient_email}')
    except Exception as e:
        print(f'Failed to send email: {e}')

if __name__ == '__main__':
    app.run(debug=True)

